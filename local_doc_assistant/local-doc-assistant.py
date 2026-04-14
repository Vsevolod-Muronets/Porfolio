#!/usr/bin/env python3
"""
Веб-интерфейс RAG-ассистента для консультирования по нормативной базе организации и редактирования заданий для договоров ГПХ.
Режим 1: вопросы по базе нормативных документов
Режим 2: загрузка DOCX-задания для договора ГПХ -> проверка -> скачивание исправленного файла
"""

import os, sys, glob, json, re, uuid, shutil, zipfile, threading
from pathlib import Path
from datetime import datetime

def check_and_install():
    missing = []
    for pkg, imp in [
        ("flask", "flask"), ("langchain", "langchain"),
        ("langchain_community", "langchain_community"),
        ("chromadb", "chromadb"), ("pypdf", "pypdf"),
        ("ollama", "ollama"), ("python-docx", "docx")
    ]:
        try:
            __import__(imp)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"\n[ОШИБКА] Не установлены пакеты: {', '.join(missing)}")
        print(f"Установите: pip install {' '.join(missing)}")
        sys.exit(1)

check_and_install()

from flask import Flask, request, jsonify, send_file, render_template_string
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import docx as python_docx

CFG = {
    "llm_model": "qwen3:8b",
    "embed_model": "qwen3-embedding:4b",
    "ollama_url": "http://localhost:11434",
    "norms_dir": "./norms",                     # нормативные документы
    "db_dir": "./chroma_db",
    "uploads_dir": "./uploads",                 # временные загруженные договоры
    "chunk_size": 1000, "chunk_overlap": 200,
    "top_k": 5, "port": 7860
}

Path(CFG["norms_dir"]).mkdir(exist_ok=True)
Path(CFG["uploads_dir"]).mkdir(exist_ok=True)

qa_chain = None
db = None
retriever = None
init_status = {"done": False, "error": None, "message": "Инициализация..."}

NORMS_PROMPT = PromptTemplate(
    template="""Ты - юридический ассистент организации, помогающий анализировать нормативные документы и договоры ГПХ.
Отвечай ТОЛЬКО на основе предоставленного контекста из нормативной базы.
Если информации недостаточно - честно скажи об этом.
Отвечай на русском языке, четко и по делу.

Контекст из нормативной базы:
{context}

Вопрос: {question}

Ответ:""",
    input_variables=["context", "question"],
)

EDITOR_PROMPT = """Ты - юридический редактор договоров ГПХ организации.
 
Тебе дано ЗАДАНИЕ для договора ГПХ (раздел, описывающий предмет и объем работ):
---
{task_text}
---
 
Нормативная база организации (требования к формулировкам):
---
{norms_context}
---
 
Твоя задача:
1. Проверь задание на соответствие нормативной базе организации.
2. Перепиши задание так, чтобы оно полностью соответствовало требованиям.
3. Сохрани суть (вид работ, сроки, результат), но исправь формулировки.
4. Если что-то противоречит нормам - исправь и кратко поясни что именно.
 
Верни ТОЛЬКО переписанный текст задания. После него - через строку "---КОММЕНТАРИЙ---" - добавь краткий список изменений (до 5 пунктов).
"""


def load_norms_db(rebuild=False):
    global db, qa_chain, retriever, init_status
    
    norms_path = Path(CFG["norms_dir"])
    all_files = (
        glob.glob(str(norms_path / "**/*.pdf"),  recursive=True) +
        glob.glob(str(norms_path / "**/*.docx"), recursive=True) +
        glob.glob(str(norms_path / "**/*.txt"),  recursive=True)
    )

    if not all_files:
        init_status["message"] = (
            f"Папка '{CFG['norms_dir']}' пуста. "
            "Положите туда нормативные документы и перезапустите сервер."
        )
        init_status["done"] = True
        return

    documents = []
    for fpath in all_files:
        fname = Path(fpath).name
        try:
            if fpath.endswith(".pdf"):
                loader = PyPDFLoader(fpath)
            elif fpath.endswith(".docx"):
                loader = Docx2txtLoader(fpath)
            else:
                loader = TextLoader(fpath, encoding="utf-8")
            docs = loader.load()
            for d in docs:
                d.metadata["source_file"] = fname
            documents.extend(docs)
        except Exception as e:
            print(f"[WARN] Пропущен {fname}: {e}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CFG["chunk_size"],
        chunk_overlap=CFG["chunk_overlap"],
    )
    chunks = splitter.split_documents(documents)

    embeddings = OllamaEmbeddings(
        model=CFG["embed_model"],
        base_url=CFG["ollama_url"],
    )

    db_path = CFG["db_dir"]
    if os.path.exists(db_path) and not rebuild:
        db = Chroma(persist_directory=db_path, embedding_function=embeddings)
        if db._collection.count() == 0:
            db = Chroma.from_documents(chunks, embedding=embeddings, persist_directory=db_path)
    else:
        if os.path.exists(db_path):
            shutil.rmtree(db_path)
        db = Chroma.from_documents(chunks, embedding=embeddings, persist_directory=db_path)

    llm = OllamaLLM(
        model=CFG["llm_model"], base_url=CFG["ollama_url"],
        temperature=0.1, num_predict=2048
    )

    retriever = db.as_retriever(search_kwargs={"k": CFG["top_k"]})

    def format_docs(docs):
        return "\n\n".join(d.page_content for d in docs)

    qa_chain = (
        {
            "context":  retriever | format_docs,
            "question": RunnablePassthrough()
        }
        | NORMS_PROMPT
        | llm
        | StrOutputParser()
    )

    init_status["done"] = True
    init_status["message"] = (
        f"Загружено {len(all_files)} файлов нормативной базы "
        f"({db._collection.count()} векторов)"
    )


def get_norms_context(query: str, k: int = 5) -> str:
    if db is None:
        return ""
    docs = db.similarity_search(query, k=k)
    return "\n\n".join(d.page_content for d in docs)

def extract_task_from_docx(filepath: str) -> tuple[str, list]:
    doc = python_docx.Document(filepath)

    TARGET_FIELDS = [
        "наименование работ/услуг", "критерии качества работ/услуг",
        "результат", "материальные носители"
    ]

    rows = []

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                label = cells[0].lower()
                if any(field in label for field in TARGET_FIELDS):
                    rows.append(f"{cells[0]}:\n{cells[1]}")

    if not rows:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                rows.append(text)

    full_text = "\n\n".join(rows)
    return full_text, []

def rewrite_docx(original_path: str, new_task_text: str, comment_text: str) -> str:
    doc = python_docx.Document(original_path)

    content_paras = [i for i, p in enumerate(doc.paragraphs) if p.text.strip()]

    if not content_paras:
        doc.add_paragraph(new_task_text)
    else:
        first_idx = content_paras[0]

        new_lines = [l.strip() for l in new_task_text.strip().split("\n") if l.strip()]

        para = doc.paragraphs[first_idx]
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_lines[0] if new_lines else ""
        else:
            para.add_run(new_lines[0] if new_lines else "")

        for idx in content_paras[1:]:
            for run in doc.paragraphs[idx].runs:
                run.text = ""

        import copy
        from docx.oxml.ns import qn
        from lxml import etree

        ref_para = doc.paragraphs[first_idx]._element
        parent   = ref_para.getparent()
        pos       = list(parent).index(ref_para)

        for i, line in enumerate(new_lines[1:], 1):
            new_p = copy.deepcopy(ref_para)
            for r in new_p.findall(f".//{qn('w:r')}"):
                new_p.remove(r)
            r_elem = etree.SubElement(new_p, qn("w:r"))
            t_elem = etree.SubElement(r_elem, qn("w:t"))
            t_elem.text = line
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            parent.insert(pos + i, new_p)

    doc.add_paragraph("")
    sep = doc.add_paragraph("─" * 50)
    sep.style = doc.styles["Normal"]

    editor_title = doc.add_paragraph("Комментарий редактора (ИИ-ассистент):")
    editor_title.runs[0].bold = True

    for line in comment_text.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    out_name = f"edited_{uuid.uuid4().hex[:8]}.docx"
    out_path = os.path.join(CFG["uploads_dir"], out_name)
    doc.save(out_path)
    return out_path


def call_llm_direct(prompt: str) -> str:
    import ollama as ol
    response = ol.chat(
        model=CFG["llm_model"],
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": 0.1, "num_predict": 3000},
    )
    return response["message"]["content"]


app = Flask(__name__)

HTML = r"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Административный ассистент</title>
<style>
  :root {
    --bg: #0f1117;
    --surface: #1a1d27;
    --border: #2a2d3e;
    --accent: #4f8ef7;
    --accent2: #7c5cbf;
    --text: #e2e4ef;
    --muted: #7b7f99;
    --green: #3ecf8e;
    --red: #f76f6f;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: var(--bg); color: var(--text); font-family: 'Segoe UI', system-ui, sans-serif;
         min-height: 100vh; display: flex; flex-direction: column; }
  header { background: var(--surface); border-bottom: 1px solid var(--border);
           padding: 16px 32px; display: flex; align-items: center; gap: 12px; }
  header h1 { font-size: 1.1rem; font-weight: 600; }
  header span { font-size: 1.5rem; }
  .status-pill { margin-left: auto; font-size: 0.75rem; padding: 4px 12px;
                 border-radius: 99px; border: 1px solid var(--border);
                 background: var(--bg); color: var(--muted); }
  .status-pill.ok  { border-color: var(--green); color: var(--green); }
  .status-pill.err { border-color: var(--red);   color: var(--red); }

  .tabs { display: flex; gap: 0; border-bottom: 1px solid var(--border); background: var(--surface); }
  .tab  { padding: 14px 28px; cursor: pointer; font-size: 0.9rem; color: var(--muted);
           border-bottom: 2px solid transparent; transition: all .2s; }
  .tab:hover  { color: var(--text); }
  .tab.active { color: var(--accent); border-color: var(--accent); }

  .panel { display: none; flex: 1; padding: 28px 32px; max-width: 900px; width: 100%; margin: 0 auto; }
  .panel.active { display: flex; flex-direction: column; gap: 16px; }

  label { font-size: 0.8rem; color: var(--muted); text-transform: uppercase; letter-spacing: .05em; }

  textarea, input[type=text] {
    width: 100%; background: var(--surface); border: 1px solid var(--border);
    border-radius: 8px; color: var(--text); padding: 12px 16px; font-size: 0.95rem;
    font-family: inherit; resize: vertical; outline: none; transition: border-color .2s;
  }
  textarea:focus, input[type=text]:focus { border-color: var(--accent); }

  .btn { padding: 10px 24px; border-radius: 8px; border: none; font-size: 0.9rem;
         cursor: pointer; font-weight: 600; transition: all .2s; }
  .btn-primary { background: var(--accent); color: #fff; }
  .btn-primary:hover { background: #3d7de8; }
  .btn-primary:disabled { background: var(--border); color: var(--muted); cursor: not-allowed; }
  .btn-success { background: var(--green); color: #0f1117; }
  .btn-success:hover { opacity: .85; }

  .answer-box { background: var(--surface); border: 1px solid var(--border);
                border-radius: 8px; padding: 16px 20px; line-height: 1.7;
                white-space: pre-wrap; font-size: 0.93rem; min-height: 80px; }
  .sources { font-size: 0.8rem; color: var(--muted); margin-top: 8px; }

  .drop-zone { border: 2px dashed var(--border); border-radius: 12px; padding: 40px;
               text-align: center; cursor: pointer; transition: all .2s; }
  .drop-zone:hover, .drop-zone.drag { border-color: var(--accent); background: rgba(79,142,247,.05); }
  .drop-zone p { color: var(--muted); font-size: 0.9rem; margin-top: 8px; }

  .progress { display: none; flex-direction: column; gap: 8px; }
  .progress.show { display: flex; }
  .progress-bar { height: 4px; background: var(--border); border-radius: 99px; overflow: hidden; }
  .progress-fill { height: 100%; width: 30%; background: linear-gradient(90deg, var(--accent), var(--accent2));
                   border-radius: 99px; animation: slide 1.5s ease-in-out infinite; }
  @keyframes slide { 0%{width:10%} 50%{width:80%} 100%{width:10%} }
  .progress-text { font-size: 0.85rem; color: var(--muted); }

  .result-block { background: var(--surface); border: 1px solid var(--green);
                  border-radius: 8px; padding: 20px; }
  .result-block h3 { color: var(--green); margin-bottom: 12px; font-size: 0.9rem; }
  .result-text { white-space: pre-wrap; font-size: 0.9rem; line-height: 1.7; }
  .comment-block { margin-top: 16px; border-top: 1px solid var(--border); padding-top: 16px; }
  .comment-block h3 { color: var(--accent2); margin-bottom: 8px; font-size: 0.85rem; }
  .row { display: flex; gap: 12px; align-items: flex-start; }
  .spinner { display:none; width:18px; height:18px; border:2px solid var(--border);
             border-top-color: var(--accent); border-radius:50%; animation: spin .6s linear infinite; }
  .spinner.show { display:inline-block; }
  @keyframes spin { to { transform: rotate(360deg); } }
</style>
</head>
<body>
<header>
  <h1>Административный ассистент</h1>
  <div class="status-pill" id="status-pill">загрузка...</div>
</header>

<div class="tabs">
  <div class="tab active" onclick="switchTab('qa')">База знаний</div>
  <div class="tab" onclick="switchTab('editor')">Редактор заданий</div>
</div>

<!-- ВКЛАДКА 1: ВОПРОСЫ К БАЗЕ -->
<div class="panel active" id="panel-qa">
  <label>Вопрос по нормативной базе организации</label>
  <textarea id="qa-input" rows="3"
    placeholder="Например: Какие требования предъявляются к описанию предмета договора ГПХ?"></textarea>
  <div class="row">
    <button class="btn btn-primary" id="qa-btn" onclick="askQuestion()">Спросить</button>
    <div class="spinner" id="qa-spinner"></div>
  </div>
  <div class="answer-box" id="qa-answer" style="display:none"></div>
  <div class="sources" id="qa-sources"></div>
</div>

<!-- ВКЛАДКА 2: РЕДАКТОР DOCX -->
<div class="panel" id="panel-editor">
  <label>Загрузите DOCX с заданием для договора ГПХ</label>
  <div class="drop-zone" id="drop-zone"
       ondragover="event.preventDefault();this.classList.add('drag')"
       ondragleave="this.classList.remove('drag')"
       ondrop="handleDrop(event)"
       onclick="document.getElementById('file-input').click()">
    <div style="font-size:2rem"></div>
    <p>Перетащите .docx файл сюда или нажмите для выбора</p>
    <p id="file-name" style="color:var(--accent);margin-top:6px"></p>
  </div>
  <input type="file" id="file-input" accept=".docx" style="display:none" onchange="handleFile(this)">

  <label>Инструкция для редактора (опционально)</label>
  <input type="text" id="editor-instruction"
    placeholder="Например: Обрати особое внимание на описание результата работ">

  <div class="row">
    <button class="btn btn-primary" id="edit-btn" onclick="editDocument()" disabled>
      Проверить и переписать
    </button>
    <div class="spinner" id="edit-spinner"></div>
  </div>

  <div class="progress" id="edit-progress">
    <div class="progress-bar"><div class="progress-fill"></div></div>
    <div class="progress-text" id="progress-text">Анализирую документ...</div>
  </div>

  <div id="edit-result" style="display:none">
    <div class="result-block">
      <h3>Переписанное задание</h3>
      <div class="result-text" id="rewritten-text"></div>
      <div class="comment-block">
        <h3>🗒 Что изменено</h3>
        <div class="result-text" id="comment-text" style="color:var(--muted)"></div>
      </div>
    </div>
    <button class="btn btn-success" onclick="downloadResult()" style="margin-top:12px">
      ⬇ Скачать исправленный DOCX
    </button>
  </div>
</div>

<script>
let uploadedFile = null;
let resultFileId = null;
const STEPS = [
  "Извлекаю текст задания...", "Ищу релевантные нормы в базе...",
  "Модель переписывает задание...", "Формирую DOCX файл..."
];

function switchTab(tab) {
  document.querySelectorAll('.tab').forEach((t, i) => {
    t.classList.toggle('active', ['qa','editor'][i] === tab);
  });
  document.querySelectorAll('.panel').forEach(p => p.classList.remove('active'));
  document.getElementById('panel-' + tab).classList.add('active');
}

async function pollStatus() {
  try {
    const r = await fetch('/api/status');
    const d = await r.json();
    const pill = document.getElementById('status-pill');
    pill.textContent = d.message;
    pill.className = 'status-pill ' + (d.error ? 'err' : d.done ? 'ok' : '');
    if (!d.done) setTimeout(pollStatus, 2000);
  } catch(e) { setTimeout(pollStatus, 3000); }
}
pollStatus();

async function askQuestion() {
  const q = document.getElementById('qa-input').value.trim();
  if (!q) return;
  const btn = document.getElementById('qa-btn');
  const spin = document.getElementById('qa-spinner');
  btn.disabled = true; spin.classList.add('show');
  document.getElementById('qa-answer').style.display = 'none';
  document.getElementById('qa-sources').textContent = '';
  try {
    const r = await fetch('/api/ask', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({question: q})
    });
    const d = await r.json();
    const box = document.getElementById('qa-answer');
    box.textContent = d.answer || d.error || 'Нет ответа';
    box.style.display = 'block';
    if (d.sources && d.sources.length) {
      document.getElementById('qa-sources').textContent =
        'Источники: ' + d.sources.join(' · ');
    }
  } catch(e) {
    document.getElementById('qa-answer').textContent = 'Ошибка: ' + e.message;
    document.getElementById('qa-answer').style.display = 'block';
  }
  btn.disabled = false; spin.classList.remove('show');
}
document.addEventListener('DOMContentLoaded', () => {
  document.getElementById('qa-input').addEventListener('keydown', e => {
    if (e.key === 'Enter' && (e.ctrlKey || e.metaKey)) askQuestion();
  });
});

function handleDrop(e) {
  e.preventDefault();
  document.getElementById('drop-zone').classList.remove('drag');
  const f = e.dataTransfer.files[0];
  if (f && f.name.endsWith('.docx')) setFile(f);
}
function handleFile(input) { if (input.files[0]) setFile(input.files[0]); }
function setFile(f) {
  uploadedFile = f;
  document.getElementById('file-name').textContent = f.name;
  document.getElementById('edit-btn').disabled = false;
  document.getElementById('edit-result').style.display = 'none';
}

async function editDocument() {
  if (!uploadedFile) return;
  const btn = document.getElementById('edit-btn');
  const spin = document.getElementById('edit-spinner');
  const prog = document.getElementById('edit-progress');
  const progText = document.getElementById('progress-text');
  btn.disabled = true; spin.classList.add('show');
  prog.classList.add('show');
  document.getElementById('edit-result').style.display = 'none';

  let step = 0;
  const stepInterval = setInterval(() => {
    if (step < STEPS.length) progText.textContent = STEPS[step++];
  }, 4000);

  const fd = new FormData();
  fd.append('file', uploadedFile);
  fd.append('instruction', document.getElementById('editor-instruction').value);

  try {
    const r = await fetch('/api/edit', { method: 'POST', body: fd });
    const d = await r.json();
    clearInterval(stepInterval);
    prog.classList.remove('show');
    if (d.error) { alert('Ошибка: ' + d.error); }
    else {
      resultFileId = d.file_id;
      document.getElementById('rewritten-text').textContent = d.rewritten;
      document.getElementById('comment-text').textContent   = d.comment;
      document.getElementById('edit-result').style.display = 'block';
    }
  } catch(e) {
    clearInterval(stepInterval);
    prog.classList.remove('show');
    alert('Ошибка: ' + e.message);
  }
  btn.disabled = false; spin.classList.remove('show');
}

function downloadResult() {
  if (resultFileId) window.location.href = '/api/download/' + resultFileId;
}
</script>
</body>
</html>
"""

@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/api/status")
def api_status():
    return jsonify(init_status)


@app.route("/api/ask", methods=["POST"])
def api_ask():
    if not init_status["done"]:
        return jsonify({"error": "База знаний еще загружается, подождите..."}), 503
    if qa_chain is None:
        return jsonify({"error": "RAG-цепочка не инициализирована. Проверьте папку norms/."}), 503

    data = request.get_json()
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"error": "Пустой вопрос"}), 400

    try:
        answer      = qa_chain.invoke(question)
        source_docs = retriever.invoke(question)
        sources     = list({
            d.metadata.get("source_file", "неизвестно")
            for d in source_docs
        })
        return jsonify({"answer": answer, "sources": sources})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/edit", methods=["POST"])
def api_edit():
    if not init_status["done"]:
        return jsonify({"error": "База знаний ещё загружается"}), 503

    if "file" not in request.files:
        return jsonify({"error": "Файл не передан"}), 400

    f = request.files["file"]
    if not f.filename.endswith(".docx"):
        return jsonify({"error": "Только .docx файлы"}), 400

    instruction = request.form.get("instruction", "")

    uid = uuid.uuid4().hex[:8]
    in_path = os.path.join(CFG["uploads_dir"], f"in_{uid}.docx")
    f.save(in_path)

    try:
        task_text, _ = extract_task_from_docx(in_path)
        if not task_text.strip():
            return jsonify({"error": "Документ пустой или не содержит текста"}), 400

        norms_context = get_norms_context(task_text)

        extra = f"\nДополнительная инструкция: {instruction}" if instruction else ""
        prompt = EDITOR_PROMPT.format(
            task_text=task_text,
            norms_context=norms_context or "(нормативная база не загружена)",
        ) + extra

        raw_response = call_llm_direct(prompt)

        if "---КОММЕНТАРИЙ---" in raw_response:
            parts = raw_response.split("---КОММЕНТАРИЙ---", 1)
            rewritten = parts[0].strip()
            comment   = parts[1].strip()
        else:
            rewritten = raw_response.strip()
            comment   = "Комментарий не сформирован."

        out_path = rewrite_docx(in_path, rewritten, comment)
        file_id  = os.path.basename(out_path)

        return jsonify({
            "rewritten": rewritten, "comment": comment, "file_id": file_id
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        if os.path.exists(in_path):
            os.remove(in_path)


@app.route("/api/download/<file_id>")
def api_download(file_id):
    safe_name = Path(file_id).name
    file_path = os.path.join(CFG["uploads_dir"], safe_name)
    if not os.path.exists(file_path):
        return jsonify({"error": "Файл не найден"}), 404
    return send_file(
        file_path, as_attachment=True,
        download_name=f"исправленное_задание_{datetime.now().strftime('%d%m%Y')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--norms", default=CFG["norms_dir"],  help="Папка с нормативными документами")
    parser.add_argument("--model", default=CFG["llm_model"],  help="LLM-модель Ollama")
    parser.add_argument("--port", default=CFG["port"], type=int)
    parser.add_argument("--rebuild", action="store_true", help="Пересоздать векторную базу")
    args = parser.parse_args()

    CFG["norms_dir"]  = args.norms
    CFG["llm_model"]  = args.model
    CFG["port"]       = args.port

    threading.Thread(
        target=lambda: load_norms_db(rebuild=args.rebuild),
        daemon=True
    ).start()

    print(f"""
Откройте в браузере: http://localhost:{CFG['port']}

Папка с нормативными документами: {CFG['norms_dir']}/
  -> Положите туда PDF/DOCX/TXT файлы и перезапустите
""")
    app.run(host="0.0.0.0", port=CFG["port"], debug=False)
